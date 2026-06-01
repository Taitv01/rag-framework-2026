"""
Document Loader
===============

Multi-format document loading with metadata extraction.

Supported formats:
- PDF (pypdf)
- DOCX (python-docx)
- HTML (BeautifulSoup)
- Markdown
- CSV
- JSON
- Plain text

Usage:
    loader = DocumentLoader()
    docs = loader.load("path/to/document.pdf")
    docs = loader.load_directory("path/to/docs/")
"""

import os
import json
import csv
from pathlib import Path
from typing import List, Optional, Dict, Any, Union
from dataclasses import dataclass, field

from langchain_core.documents import Document


@dataclass
class DocumentMetadata:
    """Metadata for loaded documents."""
    source: str
    file_type: str
    file_size: int
    created_at: Optional[str] = None
    modified_at: Optional[str] = None
    extra: Dict[str, Any] = field(default_factory=dict)


class DocumentLoader:
    """
    Multi-format document loader with metadata extraction.

    Supports PDF, DOCX, HTML, Markdown, CSV, JSON, and plain text files.
    Automatically detects file type and uses appropriate loader.

    Example:
        loader = DocumentLoader()

        # Load single file
        docs = loader.load("document.pdf")

        # Load directory
        docs = loader.load_directory("docs/")

        # Load with custom metadata
        docs = loader.load("document.pdf", metadata={"category": "research"})
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".html": "html",
        ".htm": "html",
        ".md": "markdown",
        ".markdown": "markdown",
        ".csv": "csv",
        ".json": "json",
        ".txt": "text",
        ".text": "text",
    }

    def __init__(self, encoding: str = "utf-8"):
        """
        Initialize document loader.

        Args:
            encoding: Default text encoding for text files
        """
        self.encoding = encoding

    def load(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Load a single document.

        Args:
            file_path: Path to the document
            metadata: Additional metadata to attach

        Returns:
            List of Document objects

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is not supported
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = self._get_file_type(file_path)

        # Get file metadata
        file_metadata = self._get_file_metadata(file_path)

        # Merge with custom metadata
        if metadata:
            file_metadata.update(metadata)

        # Load based on file type
        loader_map = {
            "pdf": self._load_pdf,
            "docx": self._load_docx,
            "html": self._load_html,
            "markdown": self._load_markdown,
            "csv": self._load_csv,
            "json": self._load_json,
            "text": self._load_text,
        }

        loader_func = loader_map.get(file_type)
        if not loader_func:
            raise ValueError(f"Unsupported file type: {file_type}")

        docs = loader_func(file_path, file_metadata)

        return docs

    def load_directory(
        self,
        directory: Union[str, Path],
        glob_pattern: str = "**/*",
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Load all documents from a directory.

        Args:
            directory: Path to directory
            glob_pattern: Glob pattern for file matching
            metadata: Additional metadata to attach to all documents

        Returns:
            List of Document objects from all files
        """
        directory = Path(directory)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        all_docs = []

        for file_path in directory.glob(glob_pattern):
            if file_path.is_file():
                file_ext = file_path.suffix.lower()
                if file_ext in self.SUPPORTED_EXTENSIONS:
                    try:
                        docs = self.load(file_path, metadata)
                        all_docs.extend(docs)
                    except Exception as e:
                        print(f"Warning: Failed to load {file_path}: {e}")

        return all_docs

    def _get_file_type(self, file_path: Path) -> str:
        """Detect file type from extension."""
        ext = file_path.suffix.lower()
        file_type = self.SUPPORTED_EXTENSIONS.get(ext)
        if not file_type:
            raise ValueError(
                f"Unsupported file extension: {ext}. "
                f"Supported: {list(self.SUPPORTED_EXTENSIONS.keys())}"
            )
        return file_type

    def _get_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract file metadata."""
        stat = file_path.stat()

        return {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": self._get_file_type(file_path),
            "file_size": stat.st_size,
            "file_extension": file_path.suffix.lower(),
        }

    def _load_pdf(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load PDF document."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF loading. "
                "Install it with: pip install pypdf"
            )

        reader = PdfReader(str(file_path))
        docs = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                page_metadata = metadata.copy()
                page_metadata["page_number"] = i + 1
                page_metadata["total_pages"] = len(reader.pages)

                docs.append(Document(
                    page_content=text,
                    metadata=page_metadata
                ))

        return docs if docs else [Document(page_content="", metadata=metadata)]

    def _load_docx(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load DOCX document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. "
                "Install it with: pip install python-docx"
            )

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        if not paragraphs:
            return [Document(page_content="", metadata=metadata)]

        # Combine paragraphs with proper spacing
        full_text = "\n\n".join(paragraphs)

        return [Document(page_content=full_text, metadata=metadata)]

    def _load_html(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load HTML document."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required for HTML loading. "
                "Install it with: pip install beautifulsoup4"
            )

        with open(file_path, "r", encoding=self.encoding) as f:
            soup = BeautifulSoup(f.read(), "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text(separator="\n", strip=True)

        # Extract title if available
        title = soup.find("title")
        if title:
            metadata["title"] = title.get_text(strip=True)

        return [Document(page_content=text, metadata=metadata)]

    def _load_markdown(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load Markdown document."""
        with open(file_path, "r", encoding=self.encoding) as f:
            content = f.read()

        # Extract title from first heading
        lines = content.split("\n")
        for line in lines:
            if line.startswith("# "):
                metadata["title"] = line[2:].strip()
                break

        return [Document(page_content=content, metadata=metadata)]

    def _load_csv(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load CSV document."""
        docs = []

        with open(file_path, "r", encoding=self.encoding) as f:
            reader = csv.DictReader(f)

            for i, row in enumerate(reader):
                # Convert row to text
                text_parts = [f"{key}: {value}" for key, value in row.items() if value]
                text = "\n".join(text_parts)

                row_metadata = metadata.copy()
                row_metadata["row_number"] = i + 1

                docs.append(Document(
                    page_content=text,
                    metadata=row_metadata
                ))

        return docs if docs else [Document(page_content="", metadata=metadata)]

    def _load_json(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load JSON document."""
        with open(file_path, "r", encoding=self.encoding) as f:
            data = json.load(f)

        # Convert JSON to text representation
        if isinstance(data, list):
            docs = []
            for i, item in enumerate(data):
                text = json.dumps(item, indent=2, ensure_ascii=False)
                item_metadata = metadata.copy()
                item_metadata["item_index"] = i

                docs.append(Document(
                    page_content=text,
                    metadata=item_metadata
                ))
            return docs
        else:
            text = json.dumps(data, indent=2, ensure_ascii=False)
            return [Document(page_content=text, metadata=metadata)]

    def _load_text(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load plain text document."""
        with open(file_path, "r", encoding=self.encoding) as f:
            content = f.read()

        return [Document(page_content=content, metadata=metadata)]
